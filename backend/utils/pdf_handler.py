# PDF Handler Utility Module
# File: backend/utils/pdf_handler.py
# Uses pdf2docx for PDF to Word and pypdf for compression

import fitz  # PyMuPDF
import os
import tempfile
from PIL import Image
from PyPDF2 import PdfReader, PdfWriter, PdfMerger
import io


class PDFHandler:
    """
    Comprehensive PDF manipulation handler using PyPDF2, PyMuPDF, and PIL
    Returns dict with 'success' key for all operations
    All methods handle edge cases and provide meaningful error messages
    """
    
    def __init__(self, output_dir=None):
        self.output_dir = output_dir or tempfile.gettempdir()
        self.supported_operations = [
            'merge', 'split', 'rotate', 'crop', 'compress',
            'watermark', 'extract_text', 'extract_images',
            'add_page_numbers', 'remove_pages', 'rearrange',
            'protect', 'unlock', 'pdf_to_word', 'pdf_to_images'
        ]
    
    def _validate_pdf_path(self, path, check_exists=True):
        """Validate PDF file path"""
        if not path:
            return False, "No file path provided"
        if check_exists and not os.path.exists(path):
            return False, f"File not found: {path}"
        if check_exists and os.path.getsize(path) == 0:
            return False, "File is empty"
        return True, None
    
    def _safe_open_pdf(self, path, password=None):
        """Safely open a PDF file with error handling"""
        try:
            doc = fitz.open(path)
            if doc.is_encrypted:
                if password:
                    if not doc.authenticate(password):
                        doc.close()
                        return None, "Invalid password"
                else:
                    doc.close()
                    return None, "PDF is encrypted. Password required."
            if doc.page_count == 0:
                doc.close()
                return None, "PDF has no pages"
            return doc, None
        except Exception as e:
            return None, f"Failed to open PDF: {str(e)}"
    
    def merge_pdfs(self, input_files, output_path):
        """Merge multiple PDFs into a single document"""
        try:
            # Validate inputs
            if not input_files:
                return {'success': False, 'error': 'No input files provided'}
            
            valid_files = []
            for pdf_file in input_files:
                valid, error = self._validate_pdf_path(pdf_file)
                if valid:
                    valid_files.append(pdf_file)
            
            if not valid_files:
                return {'success': False, 'error': 'No valid PDF files to merge'}
            
            if len(valid_files) < 2:
                return {'success': False, 'error': 'At least 2 valid PDF files required for merge'}
            
            merger = PdfMerger()
            
            for pdf_file in valid_files:
                try:
                    merger.append(pdf_file)
                except Exception as e:
                    # Skip problematic files but continue
                    continue
            
            merger.write(output_path)
            merger.close()
            
            if not os.path.exists(output_path) or os.path.getsize(output_path) == 0:
                return {'success': False, 'error': 'Failed to create merged PDF'}
            
            return {'success': True, 'output_path': output_path, 'files_merged': len(valid_files)}
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def split_pdf(self, input_path, output_dir, pages_str):
        """Extract specific pages from PDF"""
        try:
            valid, error = self._validate_pdf_path(input_path)
            if not valid:
                return {'success': False, 'error': error}
            
            reader = PdfReader(input_path)
            total_pages = len(reader.pages)
            
            if total_pages == 0:
                return {'success': False, 'error': 'PDF has no pages'}
            
            output_files = []
            
            # Parse pages string (e.g., "1-3" or "1,3,5" or "all")
            pages = []
            if not pages_str or pages_str.lower() == 'all':
                pages = list(range(1, total_pages + 1))
            elif '-' in pages_str:
                try:
                    parts = pages_str.split('-')
                    start = int(parts[0].strip())
                    end = int(parts[1].strip())
                    pages = list(range(start, end + 1))
                except:
                    return {'success': False, 'error': f'Invalid page range: {pages_str}'}
            elif ',' in pages_str:
                try:
                    pages = [int(p.strip()) for p in pages_str.split(',')]
                except:
                    return {'success': False, 'error': f'Invalid page list: {pages_str}'}
            else:
                try:
                    pages = [int(pages_str.strip())]
                except:
                    return {'success': False, 'error': f'Invalid page number: {pages_str}'}
            
            # Validate page numbers
            valid_pages = [p for p in pages if 1 <= p <= total_pages]
            if not valid_pages:
                return {'success': False, 'error': f'No valid pages specified. PDF has {total_pages} pages.'}
            
            os.makedirs(output_dir, exist_ok=True)
            
            for page_num in valid_pages:
                writer = PdfWriter()
                writer.add_page(reader.pages[page_num - 1])
                
                output_file = os.path.join(output_dir, f"page_{page_num}.pdf")
                with open(output_file, 'wb') as f:
                    writer.write(f)
                output_files.append(output_file)
            
            return {'success': True, 'pages': output_files, 'page_count': len(output_files)}
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def rotate_pdf(self, input_path, output_path, angle, pages=None):
        """Rotate PDF pages"""
        try:
            valid, error = self._validate_pdf_path(input_path)
            if not valid:
                return {'success': False, 'error': error}
            
            # Normalize angle to 0, 90, 180, 270
            angle = int(angle) % 360
            if angle not in [0, 90, 180, 270]:
                # Round to nearest 90
                angle = round(angle / 90) * 90 % 360
            
            doc, error = self._safe_open_pdf(input_path)
            if error:
                return {'success': False, 'error': error}
            
            total_pages = len(doc)
            
            if pages is None:
                pages = list(range(1, total_pages + 1))
            elif isinstance(pages, str):
                if pages.lower() == 'all':
                    pages = list(range(1, total_pages + 1))
                else:
                    try:
                        pages = [int(p.strip()) for p in pages.split(',')]
                    except:
                        pages = list(range(1, total_pages + 1))
            
            rotated_count = 0
            for page_num in pages:
                if 1 <= page_num <= total_pages:
                    page = doc[page_num - 1]
                    page.set_rotation(page.rotation + angle)
                    rotated_count += 1
            
            doc.save(output_path)
            doc.close()
            
            return {'success': True, 'output_path': output_path, 'pages_rotated': rotated_count}
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def compress_pdf(self, input_path, output_path, quality='medium', target_reduction=None):
        """
        Compress PDF file using pypdf (lossless) with PIL for image recompression
        
        Args:
            input_path: Path to input PDF
            output_path: Path for output PDF
            quality: 'low' (max compression), 'medium', 'high' (minimal compression)
            target_reduction: Optional target reduction percentage (10-90)
        """
        try:
            valid, error = self._validate_pdf_path(input_path)
            if not valid:
                return {'success': False, 'error': error}
            
            original_size = os.path.getsize(input_path)
            
            if original_size == 0:
                return {'success': False, 'error': 'Input file is empty'}
            
            # Import pypdf
            try:
                from pypdf import PdfWriter as PypdfWriter
            except ImportError:
                # Fall back to PyPDF2
                from PyPDF2 import PdfWriter as PypdfWriter
            
            # Determine compression settings based on quality
            if target_reduction:
                try:
                    reduction_pct = min(max(int(target_reduction), 10), 90)
                except:
                    reduction_pct = 50
                if reduction_pct >= 70:
                    img_quality = 20
                    img_scale = 0.35
                elif reduction_pct >= 50:
                    img_quality = 40
                    img_scale = 0.55
                elif reduction_pct >= 30:
                    img_quality = 60
                    img_scale = 0.75
                else:
                    img_quality = 80
                    img_scale = 0.9
            elif quality == 'low':
                img_quality = 20
                img_scale = 0.35
            elif quality == 'high':
                img_quality = 85
                img_scale = 1.0
            else:  # medium
                img_quality = 50
                img_scale = 0.6
            
            images_compressed = 0
            method_used = 'pypdf'
            
            # Method 1: Use pypdf with compress_content_streams (lossless)
            try:
                writer = PypdfWriter(clone_from=input_path)
                
                for page in writer.pages:
                    page.compress_content_streams()
                
                with open(output_path, 'wb') as f:
                    writer.write(f)
                
                compressed_size = os.path.getsize(output_path)
            except Exception as pypdf_error:
                # Fallback: simple copy first
                import shutil
                shutil.copy(input_path, output_path)
                compressed_size = original_size
            
            # Method 2: If pypdf didn't help much and we need more compression,
            # use PyMuPDF to recompress images
            if compressed_size >= original_size * 0.9 and quality != 'high':
                try:
                    doc = fitz.open(input_path)
                    
                    for page_num in range(len(doc)):
                        page = doc[page_num]
                        
                        try:
                            image_list = page.get_images(full=True)
                        except:
                            continue
                        
                        for img_info in image_list:
                            xref = img_info[0]
                            
                            try:
                                base_image = doc.extract_image(xref)
                                if not base_image:
                                    continue
                                
                                image_bytes = base_image.get("image")
                                if not image_bytes or len(image_bytes) < 3000:
                                    continue
                                
                                img = Image.open(io.BytesIO(image_bytes))
                                
                                if img.width < 30 or img.height < 30:
                                    continue
                                
                                orig_img_size = len(image_bytes)
                                
                                # Scale down if needed
                                if img_scale < 1.0:
                                    new_w = max(30, int(img.width * img_scale))
                                    new_h = max(30, int(img.height * img_scale))
                                    img = img.resize((new_w, new_h), Image.LANCZOS)
                                
                                # Convert to RGB for JPEG
                                if img.mode in ('RGBA', 'LA', 'P'):
                                    bg = Image.new('RGB', img.size, (255, 255, 255))
                                    if img.mode == 'P':
                                        img = img.convert('RGBA')
                                    if 'A' in img.mode:
                                        bg.paste(img, mask=img.split()[-1])
                                        img = bg
                                    else:
                                        img = img.convert('RGB')
                                elif img.mode != 'RGB':
                                    img = img.convert('RGB')
                                
                                # Compress to JPEG
                                buf = io.BytesIO()
                                img.save(buf, format='JPEG', quality=img_quality, optimize=True)
                                compressed_img = buf.getvalue()
                                
                                # Only update if significantly smaller
                                if len(compressed_img) < orig_img_size * 0.8:
                                    try:
                                        doc.xref_set_key(xref, "Filter", "/DCTDecode")
                                        doc.update_stream(xref, compressed_img)
                                        images_compressed += 1
                                    except:
                                        pass
                                        
                            except:
                                continue
                    
                    # Save with PyMuPDF compression
                    doc.save(
                        output_path,
                        garbage=4,
                        deflate=True,
                        deflate_images=True,
                        deflate_fonts=True,
                        clean=True
                    )
                    doc.close()
                    
                    compressed_size = os.path.getsize(output_path)
                    method_used = 'pymupdf+pypdf'
                    
                except Exception:
                    pass
            
            # If still larger than original, just copy
            if compressed_size >= original_size:
                import shutil
                shutil.copy(input_path, output_path)
                compressed_size = original_size
                reduction = 0.0
            else:
                reduction = round((1 - compressed_size / original_size) * 100, 2)
            
            return {
                'success': True,
                'output_path': output_path,
                'original_size': original_size,
                'compressed_size': compressed_size,
                'reduction': f"{max(reduction, 0)}%",
                'images_processed': images_compressed,
                'method': method_used
            }
        except Exception as e:
            return {'success': False, 'error': f'Compression failed: {str(e)}'}
    
    def add_watermark(self, input_path, output_path, text, opacity=0.3, angle=45, color=(0.5, 0.5, 0.5)):
        """Add watermark to PDF"""
        try:
            valid, error = self._validate_pdf_path(input_path)
            if not valid:
                return {'success': False, 'error': error}
            
            if not text or not text.strip():
                return {'success': False, 'error': 'Watermark text cannot be empty'}
            
            doc, error = self._safe_open_pdf(input_path)
            if error:
                return {'success': False, 'error': error}
            
            # Normalize opacity
            opacity = max(0.1, min(1.0, float(opacity)))
            
            for page in doc:
                rect = page.rect
                center = fitz.Point(rect.width / 2, rect.height / 2)
                
                # Create watermark with gray color based on opacity
                gray_value = 1 - (opacity * 0.5)
                text_color = (gray_value, gray_value, gray_value)
                
                # Use TextWriter for rotated text
                tw = fitz.TextWriter(page.rect)
                font = fitz.Font("helv")
                fontsize = min(60, rect.width / len(text) * 1.5)  # Scale font to fit
                
                # Add text centered
                tw.append(center, text, font=font, fontsize=fontsize)
                
                # Create rotation matrix
                mat = fitz.Matrix(1, 0, 0, 1, 0, 0).prerotate(angle)
                
                # Write text with rotation
                tw.write_text(page, morph=(center, mat), color=text_color)
            
            doc.save(output_path)
            doc.close()
            return {'success': True, 'output_path': output_path}
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def add_page_numbers(self, input_path, output_path, position="bottom-center", start_number=1):
        """Add page numbers to PDF"""
        try:
            valid, error = self._validate_pdf_path(input_path)
            if not valid:
                return {'success': False, 'error': error}
            
            doc, error = self._safe_open_pdf(input_path)
            if error:
                return {'success': False, 'error': error}
            
            total_pages = len(doc)
            start_number = max(1, int(start_number))
            
            for i, page in enumerate(doc):
                rect = page.rect
                page_num = start_number + i
                text = f"{page_num}"
                
                # Calculate position with margins
                margin = 30
                
                positions = {
                    "bottom-right": (rect.width - margin - 20, rect.height - margin),
                    "bottom-left": (margin, rect.height - margin),
                    "bottom-center": (rect.width / 2 - 10, rect.height - margin),
                    "top-right": (rect.width - margin - 20, margin + 10),
                    "top-left": (margin, margin + 10),
                    "top-center": (rect.width / 2 - 10, margin + 10),
                }
                
                x, y = positions.get(position, positions["bottom-center"])
                
                page.insert_text(
                    (x, y),
                    text,
                    fontsize=10,
                    color=(0, 0, 0)
                )
            
            doc.save(output_path)
            doc.close()
            return {'success': True, 'output_path': output_path, 'page_count': total_pages}
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def extract_text(self, input_path, output_path, output_format='txt'):
        """
        Extract text from PDF and save to file
        
        Args:
            input_path: Path to input PDF
            output_path: Path for output file
            output_format: 'txt' or 'docx'
        """
        try:
            valid, error = self._validate_pdf_path(input_path)
            if not valid:
                return {'success': False, 'error': error}
            
            doc, error = self._safe_open_pdf(input_path)
            if error:
                return {'success': False, 'error': error}
            
            text_parts = []
            
            for page_num, page in enumerate(doc):
                page_text = page.get_text().strip()
                if page_text:
                    text_parts.append(f"--- Page {page_num + 1} ---\n\n{page_text}")
            
            doc.close()
            
            text = "\n\n".join(text_parts).strip()
            
            if not text:
                return {'success': False, 'error': 'No text found in PDF. The PDF may contain only images.'}
            
            # Ensure output directory exists
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            if output_format == 'docx':
                try:
                    from docx import Document
                    document = Document()
                    
                    # Split text into paragraphs
                    for para in text.split('\n\n'):
                        if para.strip():
                            document.add_paragraph(para.strip())
                    
                    document.save(output_path)
                except ImportError:
                    # Fallback to txt if docx not available
                    output_path = output_path.replace('.docx', '.txt')
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(text)
            else:
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text)
            
            preview = text[:500] + "..." if len(text) > 500 else text
            
            return {
                'success': True, 
                'output_path': output_path, 
                'text_preview': preview,
                'char_count': len(text),
                'word_count': len(text.split())
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def extract_images(self, input_path, output_dir, format='jpg'):
        """Extract images from PDF"""
        try:
            valid, error = self._validate_pdf_path(input_path)
            if not valid:
                return {'success': False, 'error': error}
            
            doc, error = self._safe_open_pdf(input_path)
            if error:
                return {'success': False, 'error': error}
            
            os.makedirs(output_dir, exist_ok=True)
            
            image_paths = []
            image_count = 0
            
            for page_num, page in enumerate(doc):
                image_list = page.get_images(full=True)
                
                for img_index, img_info in enumerate(image_list):
                    xref = img_info[0]
                    
                    try:
                        base_image = doc.extract_image(xref)
                        if not base_image:
                            continue
                        
                        image_bytes = base_image.get("image")
                        if not image_bytes:
                            continue
                        
                        image_count += 1
                        
                        # Determine output format
                        ext = format.lower()
                        if ext not in ['jpg', 'jpeg', 'png']:
                            ext = 'jpg'
                        
                        image_path = os.path.join(output_dir, f"image_{image_count}.{ext}")
                        
                        # Convert and save image
                        image = Image.open(io.BytesIO(image_bytes))
                        
                        if ext in ['jpg', 'jpeg']:
                            if image.mode in ('RGBA', 'P', 'LA'):
                                background = Image.new('RGB', image.size, (255, 255, 255))
                                if image.mode == 'P':
                                    image = image.convert('RGBA')
                                if image.mode in ('RGBA', 'LA'):
                                    background.paste(image, mask=image.split()[-1])
                                    image = background
                                else:
                                    image = image.convert('RGB')
                            elif image.mode != 'RGB':
                                image = image.convert('RGB')
                            image.save(image_path, 'JPEG', quality=95)
                        else:
                            image.save(image_path, 'PNG')
                        
                        image_paths.append(image_path)
                        
                    except Exception as img_error:
                        continue
            
            doc.close()
            
            if image_count == 0:
                return {'success': False, 'error': 'No images found in PDF'}
            
            return {'success': True, 'images': image_paths, 'image_count': image_count}
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def protect_pdf(self, input_path, output_path, password):
        """Add password protection to PDF"""
        try:
            valid, error = self._validate_pdf_path(input_path)
            if not valid:
                return {'success': False, 'error': error}
            
            if not password or len(password) < 1:
                return {'success': False, 'error': 'Password cannot be empty'}
            
            doc, error = self._safe_open_pdf(input_path)
            if error:
                return {'success': False, 'error': error}
            
            doc.save(
                output_path,
                encryption=fitz.PDF_ENCRYPT_AES_256,
                owner_pw=password,
                user_pw=password,
                permissions=(
                    fitz.PDF_PERM_PRINT |
                    fitz.PDF_PERM_COPY |
                    fitz.PDF_PERM_MODIFY
                )
            )
            doc.close()
            return {'success': True, 'output_path': output_path}
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def unlock_pdf(self, input_path, output_path, password):
        """Remove password protection from PDF"""
        try:
            valid, error = self._validate_pdf_path(input_path)
            if not valid:
                return {'success': False, 'error': error}
            
            doc = fitz.open(input_path)
            
            if doc.is_encrypted:
                if not password:
                    doc.close()
                    return {'success': False, 'error': 'Password required to unlock PDF'}
                if not doc.authenticate(password):
                    doc.close()
                    return {'success': False, 'error': 'Invalid password'}
            else:
                doc.close()
                return {'success': False, 'error': 'PDF is not password protected'}
            
            doc.save(output_path)
            doc.close()
            return {'success': True, 'output_path': output_path}
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def crop_pdf(self, input_path, output_path, coordinates):
        """Crop PDF pages"""
        try:
            valid, error = self._validate_pdf_path(input_path)
            if not valid:
                return {'success': False, 'error': error}
            
            if not coordinates or len(coordinates) != 4:
                return {'success': False, 'error': 'Invalid crop coordinates. Provide [x0, y0, x1, y1]'}
            
            doc, error = self._safe_open_pdf(input_path)
            if error:
                return {'success': False, 'error': error}
            
            for page in doc:
                rect = fitz.Rect(coordinates)
                page.set_cropbox(rect)
            
            doc.save(output_path)
            doc.close()
            return {'success': True, 'output_path': output_path}
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def remove_pages(self, input_path, output_path, pages_to_remove):
        """Remove pages from PDF"""
        try:
            valid, error = self._validate_pdf_path(input_path)
            if not valid:
                return {'success': False, 'error': error}
            
            if not pages_to_remove:
                return {'success': False, 'error': 'No pages specified to remove'}
            
            doc, error = self._safe_open_pdf(input_path)
            if error:
                return {'success': False, 'error': error}
            
            total_pages = len(doc)
            
            # Validate pages
            valid_pages = [p for p in pages_to_remove if 1 <= p <= total_pages]
            if not valid_pages:
                doc.close()
                return {'success': False, 'error': f'Invalid page numbers. PDF has {total_pages} pages.'}
            
            if len(valid_pages) >= total_pages:
                doc.close()
                return {'success': False, 'error': 'Cannot remove all pages from PDF'}
            
            # Convert to 0-indexed and sort in reverse order
            pages_to_delete = sorted([p - 1 for p in valid_pages], reverse=True)
            
            for page_num in pages_to_delete:
                doc.delete_page(page_num)
            
            doc.save(output_path)
            doc.close()
            return {'success': True, 'output_path': output_path, 'pages_removed': len(valid_pages)}
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def rearrange_pages(self, input_path, output_path, new_order):
        """Rearrange PDF pages"""
        try:
            valid, error = self._validate_pdf_path(input_path)
            if not valid:
                return {'success': False, 'error': error}
            
            if not new_order:
                return {'success': False, 'error': 'No page order specified'}
            
            reader = PdfReader(input_path)
            total_pages = len(reader.pages)
            
            if total_pages == 0:
                return {'success': False, 'error': 'PDF has no pages'}
            
            writer = PdfWriter()
            pages_added = 0
            
            for page_num in new_order:
                if 1 <= page_num <= total_pages:
                    writer.add_page(reader.pages[page_num - 1])
                    pages_added += 1
            
            if pages_added == 0:
                return {'success': False, 'error': f'No valid pages in order. PDF has {total_pages} pages.'}
            
            with open(output_path, 'wb') as f:
                writer.write(f)
            
            return {'success': True, 'output_path': output_path, 'pages_reordered': pages_added}
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def extract_tables(self, input_path, output_dir, output_format='txt'):
        """
        Extract tables from PDF
        
        Args:
            input_path: Path to input PDF
            output_dir: Directory for output files
            output_format: 'txt' or 'markdown'
        """
        try:
            valid, error = self._validate_pdf_path(input_path)
            if not valid:
                return {'success': False, 'error': error}
            
            doc, error = self._safe_open_pdf(input_path)
            if error:
                return {'success': False, 'error': error}
            
            os.makedirs(output_dir, exist_ok=True)
            
            tables_found = []
            all_tables_text = ""
            
            for page_num, page in enumerate(doc):
                try:
                    tabs = page.find_tables()
                    
                    if tabs and tabs.tables:
                        for table_idx, table in enumerate(tabs.tables):
                            table_data = table.extract()
                            
                            if table_data and len(table_data) > 0:
                                table_name = f"page{page_num + 1}_table{table_idx + 1}"
                                
                                if output_format == 'markdown':
                                    md_table = self._table_to_markdown(table_data)
                                    all_tables_text += f"## Table from Page {page_num + 1}, Table {table_idx + 1}\n\n"
                                    all_tables_text += md_table + "\n\n"
                                else:
                                    txt_table = self._table_to_text(table_data)
                                    all_tables_text += f"=== Table from Page {page_num + 1}, Table {table_idx + 1} ===\n\n"
                                    all_tables_text += txt_table + "\n\n"
                                
                                tables_found.append(table_name)
                except Exception:
                    continue
            
            doc.close()
            
            if not tables_found:
                return {'success': False, 'error': 'No tables found in PDF'}
            
            ext = 'md' if output_format == 'markdown' else 'txt'
            output_path = os.path.join(output_dir, f"tables.{ext}")
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(all_tables_text)
            
            return {
                'success': True,
                'output_path': output_path,
                'tables_count': len(tables_found),
                'tables': tables_found
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _table_to_markdown(self, table_data):
        """Convert table data to markdown format"""
        if not table_data or not table_data[0]:
            return ""
        
        md_lines = []
        
        # Header row
        header = [str(cell).replace('|', '\\|') if cell else "" for cell in table_data[0]]
        md_lines.append("| " + " | ".join(header) + " |")
        md_lines.append("| " + " | ".join(["---"] * len(header)) + " |")
        
        # Data rows
        for row in table_data[1:]:
            cells = [str(cell).replace('|', '\\|') if cell else "" for cell in row]
            while len(cells) < len(header):
                cells.append("")
            md_lines.append("| " + " | ".join(cells[:len(header)]) + " |")
        
        return "\n".join(md_lines)
    
    def _table_to_text(self, table_data):
        """Convert table data to plain text format"""
        if not table_data:
            return ""
        
        txt_lines = []
        
        # Calculate column widths
        col_widths = []
        for row in table_data:
            for i, cell in enumerate(row):
                cell_str = str(cell) if cell else ""
                while len(col_widths) <= i:
                    col_widths.append(0)
                col_widths[i] = max(col_widths[i], len(cell_str))
        
        # Format rows
        for row in table_data:
            cells = []
            for i, cell in enumerate(row):
                cell_str = str(cell) if cell else ""
                if i < len(col_widths):
                    cells.append(cell_str.ljust(col_widths[i]))
                else:
                    cells.append(cell_str)
            txt_lines.append(" | ".join(cells))
        
        return "\n".join(txt_lines)
    
    def pdf_to_images(self, input_path, output_dir, format='jpeg', dpi=150, pages=None):
        """
        Convert PDF pages to images
        
        Args:
            input_path: Path to input PDF
            output_dir: Directory for output images
            format: 'jpeg' or 'png'
            dpi: Resolution (72-300)
            pages: Page numbers (string like "1", "1-3", "1,3,5") or None for all
        """
        try:
            valid, error = self._validate_pdf_path(input_path)
            if not valid:
                return {'success': False, 'error': error}
            
            doc, error = self._safe_open_pdf(input_path)
            if error:
                return {'success': False, 'error': error}
            
            os.makedirs(output_dir, exist_ok=True)
            
            total_pages = len(doc)
            
            # Parse pages string
            page_list = None
            if isinstance(pages, str) and pages:
                pages = pages.strip()
                if pages.lower() == 'all' or pages == '':
                    page_list = None
                elif '-' in pages:
                    try:
                        parts = pages.split('-')
                        start = int(parts[0].strip())
                        end = int(parts[1].strip())
                        page_list = list(range(start, end + 1))
                    except:
                        page_list = None
                elif ',' in pages:
                    try:
                        page_list = [int(p.strip()) for p in pages.split(',')]
                    except:
                        page_list = None
                else:
                    try:
                        page_list = [int(pages)]
                    except:
                        page_list = None
            
            # Validate pages
            if page_list:
                page_list = [p for p in page_list if 1 <= p <= total_pages]
                if not page_list:
                    doc.close()
                    return {'success': False, 'error': f'Invalid page numbers. PDF has {total_pages} pages.'}
            
            # Clamp DPI to reasonable range
            dpi = max(72, min(300, int(dpi)))
            
            output_paths = []
            matrix = fitz.Matrix(dpi / 72, dpi / 72)
            
            for i, page in enumerate(doc):
                page_num = i + 1
                if page_list is None or page_num in page_list:
                    pix = page.get_pixmap(matrix=matrix)
                    
                    if format.lower() in ['jpeg', 'jpg']:
                        output_path = os.path.join(output_dir, f"page_{page_num}.jpg")
                        # Convert to PIL for JPEG with quality control
                        img = Image.open(io.BytesIO(pix.tobytes("png")))
                        if img.mode != 'RGB':
                            img = img.convert('RGB')
                        img.save(output_path, 'JPEG', quality=95)
                    else:
                        output_path = os.path.join(output_dir, f"page_{page_num}.png")
                        pix.save(output_path)
                    
                    output_paths.append(output_path)
            
            doc.close()
            
            if not output_paths:
                return {'success': False, 'error': 'No pages converted'}
            
            return {
                'success': True,
                'images': output_paths,
                'page_count': len(output_paths),
                'total_pages': total_pages
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def pdf_to_word(self, input_path, output_path):
        """
        Convert PDF to Word document (.docx) using pdf2docx library
        Preserves formatting, images, and layout much better than manual extraction
        
        Args:
            input_path: Path to input PDF
            output_path: Path for output .docx file
        """
        try:
            valid, error = self._validate_pdf_path(input_path)
            if not valid:
                return {'success': False, 'error': error}
            
            # Get page count first
            pdf_doc = fitz.open(input_path)
            if pdf_doc.page_count == 0:
                pdf_doc.close()
                return {'success': False, 'error': 'PDF has no pages'}
            total_pages = pdf_doc.page_count
            pdf_doc.close()
            
            # Ensure output directory exists
            output_dir = os.path.dirname(output_path)
            if output_dir:
                os.makedirs(output_dir, exist_ok=True)
            
            # Method 1: Try pdf2docx (best quality)
            try:
                from pdf2docx import Converter
                
                cv = Converter(input_path)
                cv.convert(output_path, start=0, end=None)
                cv.close()
                
                if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                    return {
                        'success': True,
                        'output_path': output_path,
                        'page_count': total_pages,
                        'method': 'pdf2docx'
                    }
            except ImportError:
                pass  # pdf2docx not installed, try fallback
            except Exception as pdf2docx_error:
                # pdf2docx failed, try fallback
                pass
            
            # Method 2: Fallback to manual extraction with python-docx
            try:
                from docx import Document
                from docx.shared import Pt, Inches
            except ImportError:
                return {'success': False, 'error': 'Neither pdf2docx nor python-docx installed. Run: pip install pdf2docx'}
            
            pdf_doc = fitz.open(input_path)
            word_doc = Document()
            
            for page_num in range(total_pages):
                page = pdf_doc[page_num]
                
                if page_num > 0:
                    word_doc.add_page_break()
                
                try:
                    blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
                    
                    for block in blocks:
                        block_type = block.get("type", 0)
                        
                        if block_type == 0:  # Text block
                            block_text = ""
                            for line in block.get("lines", []):
                                line_text = ""
                                for span in line.get("spans", []):
                                    line_text += span.get("text", "")
                                if line_text.strip():
                                    block_text += line_text + "\n"
                            
                            if block_text.strip():
                                para = word_doc.add_paragraph(block_text.strip())
                                try:
                                    first_span = block.get("lines", [{}])[0].get("spans", [{}])[0]
                                    font_size = first_span.get("size", 11)
                                    for run in para.runs:
                                        run.font.size = Pt(min(max(font_size, 8), 28))
                                except:
                                    pass
                        
                        elif block_type == 1:  # Image block
                            try:
                                bbox = block.get("bbox")
                                if bbox:
                                    rect = fitz.Rect(bbox)
                                    mat = fitz.Matrix(2, 2)
                                    pix = page.get_pixmap(matrix=mat, clip=rect)
                                    temp_path = os.path.join(self.output_dir, f"_temp_img_{page_num}_{id(block)}.png")
                                    pix.save(temp_path)
                                    width_inches = min(6, pix.width / 144)
                                    word_doc.add_picture(temp_path, width=Inches(width_inches))
                                    try:
                                        os.remove(temp_path)
                                    except:
                                        pass
                            except:
                                pass
                                
                except Exception:
                    text = page.get_text()
                    if text.strip():
                        word_doc.add_paragraph(text.strip())
            
            pdf_doc.close()
            word_doc.save(output_path)
            
            if not os.path.exists(output_path):
                return {'success': False, 'error': 'Failed to create Word document'}
            
            return {
                'success': True,
                'output_path': output_path,
                'page_count': total_pages,
                'method': 'python-docx'
            }
            
        except Exception as e:
            return {'success': False, 'error': f'Conversion failed: {str(e)}'}


# Utility functions
def validate_pdf(file_path):
    """Check if a file is a valid PDF"""
    try:
        if not file_path or not os.path.exists(file_path):
            return False
        if os.path.getsize(file_path) == 0:
            return False
        doc = fitz.open(file_path)
        is_valid = doc.page_count > 0
        doc.close()
        return is_valid
    except:
        return False


def get_page_count(file_path):
    """Get the number of pages in a PDF"""
    try:
        if not file_path or not os.path.exists(file_path):
            return 0
        doc = fitz.open(file_path)
        count = doc.page_count
        doc.close()
        return count
    except:
        return 0
